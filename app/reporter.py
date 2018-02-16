# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_ORIENT

_default_materials_style = 'Medium List 2 Accent 5'


def make_materials_attachment(subject, data, props, style=None):
    document = Document()

    rows = data

    sec = document.sections[-1]
    sec.orientation = WD_ORIENT.LANDSCAPE
    w = sec.page_width
    sec.page_width = sec.page_height
    sec.page_height = w

    document.add_heading(subject, 0)

    p = document.add_paragraph()

    fmt = p.paragraph_format
    fmt.left_indent = Pt(20)
    tabs = fmt.tab_stops
    tabs.add_tab_stop(Pt(125))

    line = p.add_run()
    line.bold = True
    line.text = '%s\t%s' % ('Заказ:', props.get('ClientName'))
    line.add_break()

    line = p.add_run()
    line.bold = True
    line.text = '%s\t%s' % ('Файл заказа:', props.get('FileName'))

    columns = (
        ('BatchType', WD_TAB_ALIGNMENT.LEFT,    Pt(150), 'Тип партии:'),
        ('BatchQty',  WD_TAB_ALIGNMENT.CENTER,  None,    'Кол-во в партиях:'),
        ('MName',     WD_TAB_ALIGNMENT.LEFT,    Pt(400), 'Материал:'),
        ('Qty',       WD_TAB_ALIGNMENT.CENTER,  None,    'Количество:'),
    )

    def add_table(style, rows):
        total_values = [
            'Итого позиций: %s' % str(len(rows)),
            sum([item[columns[1][0]] for item in rows]),
            None,
            sum([item[columns[3][0]] for item in rows]),
        ]

        def _set_column_value(cell, text, width=None, alignment=WD_TAB_ALIGNMENT.LEFT, bold=False):
            p = cell.paragraphs[0]
            r = p.add_run()
            r.bold = bold
            r.text = text
            fmt = p.paragraph_format
            fmt.alignment = alignment
            fmt.space_before = Pt(5)
            fmt.space_after = Pt(5)

            if width is not None:
                cell.width = width
        
        table = document.add_table(rows=1, cols=4)
        table.style = style

        headers = table.rows[0].cells
        for n, column in enumerate(columns):
            _set_column_value(headers[n], columns[n][3], 
                              width=column[2], 
                              alignment=column[1], 
                              bold=True)

        for item in rows:
            cells = table.add_row().cells
            for n, column in enumerate(columns):
                _set_column_value(cells[n], str(item[column[0]]), 
                                  alignment=column[1], 
                                  bold=False)

        totals = table.add_row().cells
        for n, column in enumerate(columns):
            _set_column_value(totals[n], str(total_values[n] or ''), 
                              alignment=column[1], 
                              bold=True)

    add_table(document.styles[style or _default_materials_style], rows)

    #document.save('w:/apps/perso/x.doc')

    return document
